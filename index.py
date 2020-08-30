import datetime
import os
import platform
import random
import sys
from time import sleep
import sqlite3
import numpy as np
import pandas as pd
from PyQt5.QtCore import QPropertyAnimation, QRect, Qt, QStringListModel, QFileInfo
from PyQt5.QtWidgets import QApplication, QWidget, QInputDialog, QHeaderView, QMenu, QAction, QCompleter, \
    QTableWidgetItem
from PyQt5.QtGui import QIcon, QPixmap, QIntValidator
from PyQt5 import uic, QtWidgets, QtCore, QtGui
from time import gmtime, strftime
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Cm

from plyer import notification

# main_ui, _ = uic.loadUiType(os.path.join(os.path.dirname(__file__), "src/ui/main.ui"))
#
# class Main(QWidget, main_ui):
#     def __init__(self):
#         super(Main, self).__init__()
#         QWidget.__init__(self)
#
#         # self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
#         self.setupUi(self)
#         self.menu_icon.setPixmap(QPixmap('src/icons/menu.png'))
#         self.menu_icon.setScaledContents(True)
from pynput import keyboard
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4, A5
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph

DB = 'src/db.db'
STORE_NAME = 'PARA-ELECTRO'

DESKTOP = os.path.join(os.path.join(os.environ['USERPROFILE']),'Desktop') if platform.system() == 'Windows' else os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop')
class Main(QtWidgets.QWidget):

    def __init__(self):
        super(Main, self).__init__()
        uic.loadUi(os.path.join(os.getcwd(), 'src/ui/main.ui'), self)
        self.refresh()
        # self.menu_frame.setStyleSheet(':hover{background:#e5f1fb}')
        # self.setWindowFlags(QtCore.Qt.FramelessWindowHint)
        # self.sideBar.setFixedWidth(68)
        # self.resize_event()

        # self.anim = QPropertyAnimation(self.sideBar, b"geometry")
        # self.anim.setDuration(200)
        # self.anim.setStartValue(QRect(1, 1, 1, self.height()))
        # self.anim.setEndValue(QRect(1, 1, 68, self.height()))
        # self.anim.start()

        conn = sqlite3.connect(DB)
        curs = conn.cursor()
        curs.execute("""
            CREATE TABLE IF NOT EXISTS products (
            id INTEGER PRIMARY KEY AUTOINCREMENT, 
            name TEXT, 
            code TEXT, 
            categorie TEXT, 
            qt INTEGER, 
            price TEXT,
            sell_price TEXT);
        """)
        curs.execute("""
            CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT, 
            date_time TEXT, 
            product TEXT, 
            qt INTEGER, 
            total_price TEXT, 
            person TEXT,
            paid TEXT,
            invoice TEXT,
            operation TEXT,
            rass_lmal INTEGER
            );
        """)
        curs.execute("""
            CREATE TABLE IF NOT EXISTS wishlist (
            id INTEGER PRIMARY KEY AUTOINCREMENT, 
            name TEXT, 
            note TEXT
            );
        """)
        # curs.execute("""
        #     CREATE TABLE IF NOT EXISTS debt (
        #     id INTEGER PRIMARY KEY AUTOINCREMENT,
        #     date_time TEXT,
        #     person TEXT,
        #     product TEXT,
        #     qt INTEGER,
        #     total_price TEXT,
        #     operation TEXT,
        #     state TEXT);
        # """)
        # curs.execute("""
        #     CREATE TABLE IF NOT EXISTS factures (
        #     id INTEGER PRIMARY KEY AUTOINCREMENT,
        #     date_time TEXT,
        #     persons TEXT,
        #     products TEXT,
        #     qt INTEGER,
        #     total_price TEXT,
        #     operation TEXT);
        # """)
        conn.commit()
        conn.close()
        self.setWindowIcon(QtGui.QIcon('src/icons/logo.png'))
        self.setGeometry(QRect(200, 100, 1242, 651))

        style = '''
            
            #sideBar{
	            border-right:1px solid black;
	            background-color:#EDEDED;
            }
            
            
            #menu_icon:hover {
                /*border:2px groove #e0e1dd;*/
                border:1px solid grey;
                background-color:#F5F5F5;
            }
            
            #home_frame:hover {
                border:1px solid grey;
            }
            
            #sell_frame:hover {
                
                border:1px solid grey;
            }
            
           #buy_frame:hover {
                border:1px solid grey;
            }
            
            
            #return_frame:hover {
                
                border:1px solid grey;
            }
            
            #debt_frame:hover {
                border:1px solid grey;
            }
            
            #statistics_frame:hover {
                border:1px solid grey;
            }
            
            #settings_frame:hover {
                border:1px solid grey;
            }
            
            
            #sell_frame{border-top:1px solid grey; border-bottom:1px solid grey;}
            #return_frame{border-top:1px solid grey; border-bottom:1px solid grey;}
            #statistics_frame{border-top:1px solid grey; }
            
            



        '''

        style1 = '''
            
            #sideBar{
	            border-right:1px solid black;
	            background-color:#00BFA5;
            }
            
            
            #sell_frame{border-top:1px solid #fff; border-bottom:1px solid #fff;}
            #return_frame{border-top:1px solid #fff; border-bottom:1px solid #fff;}
            #statistics_frame{border-top:1px solid #fff; }
            #logout_frame{border-bottom:1px solid #fff; }
            
            #menu_icon:hover {
                /*border:2px groove #e0e1dd;*/
                border:1px solid grey;
                background-color:#16e2c6;
            }
            
            #home_frame:hover {
                border:1px solid grey;
            }
            
            #sell_frame:hover {
                
                border:1px solid grey;
            }
            
           #buy_frame:hover {
                border:1px solid grey;
            }
            
            
            #return_frame:hover {
                
                border:1px solid grey;
            }
            
            #debt_frame:hover {
                border:1px solid grey;
            }
            
            #statistics_frame:hover {
                border:1px solid grey;
            }
            #logout_frame:hover {
                border:1px solid grey;
            }
            
            #settings_frame:hover {
                border:1px solid grey;
            }
            
        '''
        # self.side_bg = '#EDEDED'
        self.side_bg = '#00BFA5'
        # self.hover_color = '#F5F5F5'
        self.hover_color = '#16e2c6'
        self.setStyleSheet(style1)

        self.logo.setPixmap(QPixmap('src/icons/logo.png'))
        self.logo.setScaledContents(True)



        self.menu_icon.setPixmap(QPixmap('src/icons/menu.png'))
        self.menu_icon.setScaledContents(True)


        # self.buy_icon.setPixmap(QPixmap('src/icons/buy.png'))
        # self.buy_icon.setScaledContents(True)

        self.home_icon.setPixmap(QPixmap('src/icons/home.png'))
        self.home_icon.setScaledContents(True)


        # self.return_icon.setPixmap(QPixmap('src/icons/return.png'))
        # self.return_icon.setScaledContents(True)

        self.debt_icon.setPixmap(QPixmap('src/icons/debt.png'))
        self.debt_icon.setScaledContents(True)

        self.sell_icon.setPixmap(QPixmap('src/icons/buy.png'))
        self.sell_icon.setScaledContents(True)

        self.statistics_icon.setPixmap(QPixmap('src/icons/statics.png'))
        self.statistics_icon.setScaledContents(True)

        self.logout_icon.setPixmap(QPixmap('src/icons/logout.png'))
        self.logout_icon.setScaledContents(True)

        self.settings_icon.setPixmap(QPixmap('src/icons/settings.png'))
        self.settings_icon.setScaledContents(True)

        # self.frame.setStyleSheet('box-shadow: -1px 0px 10px 0px;')
        # self.menu_icon.mousePressEvent = lambda : self.sideBar.setFixedWidth(212)
        self.menu_icon.mousePressEvent = self.open_menu
        self.m = 0
        self.home = Home()
        self.content.addWidget(self.home)
        self.setContentsMargins(0, 0, 0, 0)

        self.sideBar.installEventFilter(self)
        self.home_frame.installEventFilter(self)
        self.sell_frame.installEventFilter(self)
        # self.buy_frame.installEventFilter(self)
        # self.return_frame.installEventFilter(self)
        self.debt_frame.installEventFilter(self)
        self.statistics_frame.installEventFilter(self)
        self.logout_frame.installEventFilter(self)
        self.settings_frame.installEventFilter(self)


    def refresh(self):
        self.title_label.setText(f"PARA-ELECTRO : {' - '.join(str(datetime.datetime.today().date()).split('-')[::-1])}")

    def eventFilter(self, source, event):
        if (event.type() == QtCore.QEvent.FocusOut and source is self.sideBar):
            # print('eventFilter: Side bar focused out')
            if self.m:
                self.open_menu(True)

        elif (event.type() == QtCore.QEvent.HoverEnter and source is self.sell_frame):
            self.sell_frame.setStyleSheet(f'background-color: {self.hover_color}')

        elif (event.type() == QtCore.QEvent.HoverLeave and source is self.sell_frame):
            self.sell_frame.setStyleSheet(f'background-color: {self.side_bg};')

        elif (event.type() == QtCore.QEvent.HoverEnter and source is self.home_frame):
            self.home_frame.setStyleSheet(f'background-color: {self.hover_color};')

        elif (event.type() == QtCore.QEvent.HoverLeave and source is self.home_frame):
            self.home_frame.setStyleSheet(f'background-color: {self.side_bg};')

        # elif (event.type() == QtCore.QEvent.HoverEnter and source is self.buy_frame):
        #     self.buy_frame.setStyleSheet(f'background-color: {self.hover_color};')

        # elif (event.type() == QtCore.QEvent.HoverLeave and source is self.buy_frame):
        #     self.buy_frame.setStyleSheet(f'background-color: {self.side_bg};')

        # elif (event.type() == QtCore.QEvent.HoverEnter and source is self.return_frame):
        #     self.return_frame.setStyleSheet(f'background-color: {self.hover_color};')

        # elif (event.type() == QtCore.QEvent.HoverLeave and source is self.return_frame):
        #     self.return_frame.setStyleSheet(f'background-color: {self.side_bg};')

        elif (event.type() == QtCore.QEvent.HoverEnter and source is self.debt_frame):
            self.debt_frame.setStyleSheet(f'background-color: {self.hover_color};')

        elif (event.type() == QtCore.QEvent.HoverLeave and source is self.debt_frame):
            self.debt_frame.setStyleSheet(f'background-color: {self.side_bg};')

        elif (event.type() == QtCore.QEvent.HoverEnter and source is self.statistics_frame):
            self.statistics_frame.setStyleSheet(f'background-color: {self.hover_color};')

        elif (event.type() == QtCore.QEvent.HoverLeave and source is self.statistics_frame):
            self.statistics_frame.setStyleSheet(f'background-color: {self.side_bg};')

        elif (event.type() == QtCore.QEvent.HoverEnter and source is self.logout_frame):
            self.logout_frame.setStyleSheet(f'background-color: {self.hover_color};')

        elif (event.type() == QtCore.QEvent.HoverLeave and source is self.logout_frame):
            self.logout_frame.setStyleSheet(f'background-color: {self.side_bg};')

        elif (event.type() == QtCore.QEvent.HoverEnter and source is self.settings_frame):
            self.settings_frame.setStyleSheet(f'background-color: {self.hover_color};')

        elif (event.type() == QtCore.QEvent.HoverLeave and source is self.settings_frame):
            self.settings_frame.setStyleSheet(f'background-color: {self.side_bg};')






        elif (event.type() == QtCore.QEvent.MouseButtonPress and source is self.home_frame):
            self.change_widget(Home())

        elif (event.type() == QtCore.QEvent.MouseButtonPress and source is self.sell_frame):
            self.change_widget(Sell())

        # elif (event.type() == QtCore.QEvent.MouseButtonPress and source is self.buy_frame):
        #     self.change_widget(Buy())

        elif (event.type() == QtCore.QEvent.MouseButtonPress and source is self.debt_frame):
            self.dept = Debt()
            self.dept.show()



        return super(Main, self).eventFilter(source, event)

    #     self.sideBar.installEventFilter(self)
    #
    # def eventFilter(self, obj, event):
    #     if obj == self.sideBar and event.type() == QtCore.QEvent.HoverEnter:
    #         self.onHovered()
    #     return super(Main, self).eventFilter(obj, event)
    #
    # def onHovered(self):
    #     print("hovered")
    #     self.sideBar.setStyleSheet('background-color: #fff; color: #000;')

    def change_widget(self, widget):
        for i in reversed(range(self.content.count())):
            self.content.itemAt(i).widget().setParent(None)
        self.content.addWidget(widget)


    def resizeEvent(self, event):
        self.m = 0
        print(f'{self.width()} x {self.height()}')
        self.sideBar.setGeometry(QRect(0, 0, 68, self.height() ))
        self.title_frame.setGeometry(QRect(self.sideBar.width()-1, 0, self.width()-self.sideBar.width(), 70 ))
        self.frame.setGeometry(QRect(self.sideBar.width()-1, self.title_frame.height(), self.width()-self.sideBar.width()-2, self.height()-self.title_frame.height()))

    def open_menu(self, event):
        duration, w0, w1 = 200, 68, 270
        if self.m:
            self.anim = QPropertyAnimation(self.sideBar, b"geometry")
            self.anim.setDuration(duration)
            self.anim.setStartValue(QRect(0, 0, w1, self.height()))
            self.anim.setEndValue(QRect(0, 0, w0, self.height()))
            self.anim.start()
            self.m = 0
        else:
            self.anim = QPropertyAnimation(self.sideBar, b"geometry")
            self.anim.setDuration(duration)
            self.anim.setStartValue(QRect(0, 0, w0, self.height()))
            self.anim.setEndValue(QRect(0, 0, w1, self.height()))
            self.anim.start()
            self.m = 1

    # def debt(self):
    #     text, ok = QInputDialog.getText(self, 'Money Operations', 'Enter the Amount :')
    #     if ok:
    #         print(text)


class Home(QtWidgets.QFrame):
    def __init__(self):
        super(Home, self).__init__()
        uic.loadUi(os.path.join(os.getcwd(), 'src/ui/home.ui'), self)

        # self.search_btn.setPixmap(QPixmap('src/icons/search.png'))
        # self.search_btn.setScaledContents(True)
        # self.search_btn.installEventFilter(self)
        self.search_txt.installEventFilter(self)
        # self.search_txt.setFixedWidth(0)
        # self.history_table.itemSelectionChanged.connect(self.table_select_event)
        # self.table_header = ['id', 'Date/Time', 'Product', 'Quantity', 'Total Price', 'Operation', 'Person']
        # self.history_table.setColumnCount(len(self.table_header))
        # self.history_table.setHorizontalHeaderLabels(self.table_header)
        # self.history_table.resizeColumnsToContents()
        # # self.history_table.horizontalHeader().setSectionResizeMode(self.table_header.index(self.table_header[-1]), QHeaderView.Stretch)
        # for i in range(self.history_table.columnCount()):
        #     self.history_table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)
        self.refresh()


        self.comboBox.currentTextChanged.connect(self.refresh)
        self.search_txt.textChanged.connect(self.searching)

    def searching(self):
        if self.search_txt.text():
            [self.home_table.removeRow(0) for i in range(self.home_table.rowCount())]
            con = sqlite3.connect(DB)
            key = self.search_txt.text()
            data = []
            if self.comboBox.currentIndex() == 0:
                data = [[c for c in row] for row in con.execute(f'select id, date_time, product, qt, person, paid, operation, total_price, invoice from history where date_time like "%{key}%" or product like "%{key}%" or person like "%{key}%" or invoice like "%{key}%" or operation like "%{key}%"').fetchall()]

            elif self.comboBox.currentIndex() == 1:
                data = [[c for c in row] for row in con.execute(f'select * from products where name like "%{key}%" or code like "%{key}%" or categorie like "%{key}%"').fetchall()]

            if data:
                for r in data:
                    self.home_table.insertRow(0)
                    for c in range(len(r)):
                        self.home_table.setItem(0, c, QTableWidgetItem(str(r[c])))
            else:
                [self.home_table.removeRow(0) for i in range(self.home_table.rowCount())]



        else:
            self.refresh()

    def refresh(self):
        con = sqlite3.connect(DB)
        cur = con.cursor()
        date = str(datetime.datetime.today().date())

        today_earn = 0
        month_earn = 0
        total_item_bought_counter = 0
        today_item_bought_counter = 0
        month_item_bought_counter = 0
        total_selling = 0
        month_selling = 0
        today_selling = 0
        total_earning = 0
        mount_client_have_to_pay = 0
        clients_gotta_pay = ['lkridi li katssal\n']
        mount_u_have_to_pay = 0
        suppliers_gotta_get_paid = ['lkridi li 3liiik \n']
        history = [list(i) for i in con.execute('SELECT id, date_time, product, qt, person, paid, operation, total_price, invoice, rass_lmal FROM history order by date_time asc').fetchall()]
        for i in history:

            # print(f"{str(i[1]).split('-')[0]} == {date.split('-')[2]} > {str(i[1]).split('-')[0] == date.split('-')[2]}")
            # print(f"{str(i[1]).split('-')[1]} == {date.split('-')[1]} > {str(i[1]).split('-')[1] == date.split('-')[1]}")
            # print(f"{str(i[1]).split('-')[2].split(' ')[0]} == {date.split('-')[0]} > {str(i[1]).split('-')[2].split(' ')[0] == date.split('-')[0]}")
            # print(f"{i[6]} == 'sell' > {i[6] == 'sell'}")
            # print(f"{str(i[1]).split('-')[1]} == {date.split('-')[1]} > {str(i[1]).split('-')[1] == date.split('-')[1]}")
            # print(f"{str(i[1]).split('-')[2].split(' ')[0]} == {date.split('-')[0]} > {str(i[1]).split('-')[2].split(' ')[0] == date.split('-')[0]}")
            # print(f"{i[6]} == 'sell' > {i[6] == 'sell'}")

            if str(i[1]).split('-')[0] == date.split('-')[2] and str(i[1]).split('-')[1] == date.split('-')[1] and str(i[1]).split('-')[2].split(' ')[0] == date.split('-')[0] and i[6] == 'sell':  # for today
                today_earn += int(i[7]) - int(i[9])
                today_selling += int(i[7])
                today_item_bought_counter += int(i[3])

            if str(i[1]).split('-')[1] == date.split('-')[1] and str(i[1]).split('-')[2].split(' ')[0] == date.split('-')[0] and i[6] == 'sell':  # for this month
                month_earn += int(i[7]) - int(i[9])
                month_selling += int(i[7])
                month_item_bought_counter += int(i[3])

            if i[6] == 'sell':
                total_earning += int(i[7]) - int(i[9])
                total_selling += int(i[7])
                total_item_bought_counter += int(i[3])
                if str(i[5]).lower() == 'no':
                    mount_client_have_to_pay += int(i[7])
                    clients_gotta_pay.append(f'''+ {i[4]} : {i[7]}DH\n     - Product : {i[2]}\n     - Operation Id : {i[0]}\n''')
            else:
                if str(i[5]).lower() == 'no':
                    mount_u_have_to_pay += int(i[7])
                    suppliers_gotta_get_paid.append(f'''+ {i[4]} : {i[7]}DH\n     - Product : {i[2]}\n     - Operation Id : {i[0]}\n''')


        self.counter1.setText(f'{round(month_earn/2, 2)}(<font color=green>+{round(today_earn/2, 1)}</font>) DH')
        self.counter1.setToolTip(f'''
            #######  Earning #######\n
            + Total Earning : {total_earning}DH from {total_selling} DH\n
                - this Month : {month_earn}DH from {month_selling} DH\n
                - Today : {today_earn}DH from {today_selling} DH\n
            + Total Sold items : {total_item_bought_counter}\n
                - this Month : {month_item_bought_counter}\n
                - Today : {today_item_bought_counter}
        ''')
        self.counter2.setText(f'<font color=green>{mount_client_have_to_pay} DH</font>')
        self.counter2.setToolTip(f'''{"".join(clients_gotta_pay)}''')

        self.counter3.setText(f'<font color=green>{mount_u_have_to_pay} DH</font>')
        self.counter3.setToolTip(f'''{"".join(suppliers_gotta_get_paid)}''')

        producs_out_of_stock = [f' + {i[0]}({i[1]})' for i in cur.execute('select name, code from products where qt = 0').fetchall()]
        self.counter4.setText(str(len(producs_out_of_stock)))
        producs_out_of_stock.insert(0, 'Products Out Of Stock')
        self.counter4.setToolTip('\n'.join(producs_out_of_stock))
        wish_list = [f'  + {i[0]} ({i[1]})' for i in cur.execute('select name, note from wishlist order by id desc')]
        self.counter5.setText(str(len(wish_list)))
        wish_list.insert(0, 'Products to Buy : ')
        self.counter5.setToolTip('\n'.join(wish_list))
        products = [list(i) for i in cur.execute('select * from products order by qt desc').fetchall()]



        self.home_table.clear()
        for _ in range(self.home_table.rowCount()):
            self.home_table.removeRow(0)

        if self.comboBox.currentIndex() == 0:  # for histoty table
            self.search_txt.setEnabled(True)
            head = 'Id Date Product Quantity Buyer/Seller Paid Operation Total Invoice'.split(' ')
            self.add.setFixedWidth(0)
            self.home_table.setColumnCount(len(head))
            self.home_table.setHorizontalHeaderLabels(head)
            self.home_table.resizeColumnsToContents()
            # self.to_buy_table.horizontalHeader().setSectionResizeMode(self.table_header.index(self.table_header[-1]), QHeaderView.Stretch)
            for i in range(len(head)):
                self.home_table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)
            history_data = [i[:-1] for i in history]
            for row in range(len(history_data)):
                self.home_table.insertRow(0)
                for col in range(len(head)):
                    self.home_table.setItem(0, col, QTableWidgetItem(str(history_data[row][col])))

        elif self.comboBox.currentIndex() == 1:
            self.search_txt.setEnabled(True)
            self.add.setFixedWidth(85)
            products_header = [str(i[0]).replace('_', ' ').replace('qt', 'Quantity').capitalize() for i in con.execute('SELECT name FROM PRAGMA_TABLE_INFO("products");').fetchall()]
            self.home_table.setColumnCount(len(products_header))
            self.home_table.setHorizontalHeaderLabels(products_header)
            self.home_table.resizeColumnsToContents()
            # self.to_buy_table.horizontalHeader().setSectionResizeMode(self.table_header.index(self.table_header[-1]), QHeaderView.Stretch)
            for i in range(len(products_header)):
                self.home_table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)

            for row in range(len(products)):
                self.home_table.insertRow(0)
                for col in range(len(products_header)):
                    self.home_table.setItem(0, col, QTableWidgetItem(str(products[row][col])))
        else:
            self.search_txt.setEnabled(False)
        # ll = []
        # for r in range(self.home_table.rowCount()):
        #     row = []
        #     for c in range(self.home_table.columnCount()):
        #         row.append(self.home_table.item(r, c).text())
        #     ll.append(row)
        # # self.table_content.reverse()
        # self.table_content = pd.DataFrame(ll, columns=range(len(ll[0][0])))
        # self.table_content.reset_index(drop=True, inplace=True)
        con.close()


    def table_select_event(self):
        items = self.history_table.selectedItems()
        if items:
            # print([str(i.text()) for i in items])
            self.return_btn.setEnabled(True)
        else:
            self.return_btn.setEnabled(False)

    # def eventFilter(self, source, event):
    #     if source is self.search_btn:
    #
    #         if event.type() == QtCore.QEvent.MouseButtonPress:
    #             # self.search_txt.setText('')
    #             self.search_txt.setFixedWidth(311)
    #         elif event.type() == QtCore.QEvent.HoverEnter:
    #             self.search_btn.setStyleSheet('border:1px solid grey;')
    #         elif event.type() == QtCore.QEvent.HoverLeave :
    #             self.search_btn.setStyleSheet('border:0px solid grey;')
    #
    #     return super(Home, self).eventFilter(source, event)



class Sell(QtWidgets.QFrame):
    def __init__(self):
        super(Sell, self).__init__()
        uic.loadUi(os.path.join(os.getcwd(), 'src/ui/sell.ui'), self)

        self.product_name.installEventFilter(self)

        self.completer = QCompleter()
        self.model = QStringListModel()
        self.completer.setModel(self.model)
        # # self.product_name.textChanged.connect(self.product_txt_changing)
        # self.product_txt_changing()
        self.seller.installEventFilter(self)
        self.seller.setCompleter(self.completer)
        self.product_name.currentTextChanged.connect(self.product_choosen)
        self.qt_.textChanged.connect(self.qt_changing)
        self.qt_.installEventFilter(self)

        self.qt_.setValidator(QIntValidator())
        self.price.setValidator(QIntValidator())
        self.price.installEventFilter(self)
        # conn.close()
        self.qt_value = 0
        self.add.clicked.connect(self.add_row)

        self.to_buy_table.itemSelectionChanged.connect(self.table_select_event)
        self.table_header = ['Product', 'Code', 'Categorie', 'Quantity', 'Person', 'Price', 'Paid', 'Total']
        self.to_buy_table.setColumnCount(len(self.table_header))
        self.to_buy_table.setHorizontalHeaderLabels(self.table_header)
        self.to_buy_table.resizeColumnsToContents()
        # self.to_buy_table.horizontalHeader().setSectionResizeMode(self.table_header.index(self.table_header[-1]), QHeaderView.Stretch)
        for i in range(len(self.table_header)):
            self.to_buy_table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)

        self.clear.setEnabled(False)
        self.clear.clicked.connect(self.del_row)
        self.buy.clicked.connect(self.save)
        self.f_date.setText('-'.join(str(datetime.datetime.today().date()).split('-')[::-1]))
        [self.to_buy_table.removeRow(0) for _ in range(self.to_buy_table.rowCount())]

        self.refresh()
        # with keyboard.GlobalHotKeys({
        #     '<ctrl>+<alt>+h': self.save_buy()}) as h:
        #     h.join()

        # self.operation_type.stateChanged.connect(self.change_operation_type)
        self.categorie_.setReadOnly(True)

        self.qt_bought = {}



    # def change_operation_type(self):
    #     if self.operation_type.isChecked():
    #         self.categorie_.setReadOnly(False)
    #         self.seller.setPlaceholderText('Seller Name *')
    #     else:
    #         self.categorie_.setReadOnly(True)
    #         self.seller.setPlaceholderText('Buyer Name (Optional)')


    def eventFilter(self, obj, event):
        if event.type() == QtCore.QEvent.KeyPress :
            if event.key() == QtCore.Qt.Key_Return :
                if obj is self.qt_ or obj is self.seller or obj is self.price:
                    self.add_row()

        return super().eventFilter(obj, event)



    def save(self):
        if self.to_buy_table.rowCount():
            # print(self.to_buy_table.rowCount())
            # print(self.to_buy_table.columnCount())
            bought = []
            to_update = []
            to_add = []
            conn = sqlite3.connect(DB)
            curs = conn.cursor()
            exists_codes = [str(i[0]) for i in curs.execute('select code from products').fetchall()]


            exists_invoices_ids = [str(i[0]) for i in curs.execute('SELECT invoice FROM history').fetchall()]
            pp = 'a b c d e f g h i j k l m n o p q r s t u v w x y z A B C D E F G H J L M N P Q R S T U V W Y X Z 0 1 2 3 4 5 6 7 8 9'
            def get_facture_id() -> str:
                id = ''.join(random.choices(pp.split(), k=10))
                if id in exists_invoices_ids:
                    get_facture_id()
                return id
            facture_id = get_facture_id()

            dt = []
            for r in range(self.to_buy_table.rowCount()):
                tmp = []
                for c in range(self.to_buy_table.columnCount()):
                    tmp.append(str(self.to_buy_table.item(r, c).text()))
                dt.append(tmp)

            buy_time =f'{self.f_date.text()} {str(strftime(f"%H:%M:%S", gmtime()))}' if self.f_date.text() else str(strftime(f"%Y-%m-%d %H:%M:%S", gmtime()))


            for row in dt:
                curs.execute(f'''
                                    UPDATE products SET qt = qt - {int(row[3])} where code like "{row[1]}"
                                    ''')
                curs.execute(
                    f'''INSERT INTO history (date_time, product, qt, total_price, person, paid, invoice, operation, rass_lmal) values(
                                                                        "{buy_time}", 
                                                                        "{row[0]}({row[1]})", 
                                                                        "{row[3]}", 
                                                                        "{row[-1].split(' ')[0]}",
                                                                        "{row[4]}",
                                                                        "{"No" if row[6][0].upper() == "N" else buy_time}",
                                                                        "{facture_id}",
                                                                        "sell",
                                                                        {int(curs.execute(f'select price from products where code like "{row[1]}"').fetchone()[0]) * int(row[3])}
                                                                        )''')
                conn.commit()



            # for i in to_add:
            #     curs.execute(f'''INSERT INTO products (name, code, categorie, qt, price) values("{i[0]}", "{i[1]}", "{i[2]}", {int(i[3])}, "{i[5]}")''')
            #     add_history(i)
            #
            # for i in to_update:
            #     curs.execute(f'''
            #          UPDATE products SET qt = qt + {int(i[3])} where code like "{i[1]}"
            #     ''')
            #     add_history(i)

            # tt_price = 0
            # for i in range(self.to_buy_table.rowCount()):
            #     tt_price += int(self.to_buy_table.item(i, 6).text())

            # self.to_buy_table.setHorizontalHeaderLabels(self.table_header)

            if self.print_facture.isChecked():
                data = [[i.replace('_', ' ') for i in 'Product Categorie Buyer is_Paid Quantity Total'.split(' ')]]
                for i in dt:
                    data.append([f'{i[0]}({i[1]})', i[2], i[4], i[6], str(i[3]), i[-1]])







                # for row in dt:
                #     rr = [f'{str(self.to_buy_table.item(row, 0).text())}({str(self.to_buy_table.item(row, 1).text())})']
                #     for col in [2, 4, 6, 3, 7]:
                #         rr.append(str(self.to_buy_table.item(row, col).text()))
                #     qts.append(int(str(self.to_buy_table.item(row, 3).text())))
                #     prices.append(int(str(self.to_buy_table.item(row, 7).text()).split(' ')[0]))
                #     rows.append(rr)

                # rows.append([' ', ' ', ' ', ' ', sum(qts), f'{sum(prices)} DH'])

                if filename := QtWidgets.QFileDialog.getSaveFileName(caption='Print Facture', filter="PDF (*.pdf )", directory=os.path.join(DESKTOP, f'invoice_{facture_id}'))[0]:
                    if not QFileInfo(filename).suffix():
                        filename += '.pdf'

                    # doc = SimpleDocTemplate(filename=filename, pagesize=letter, bottomMargin=18)
                    # table = Table(data=rows)
                    doc = SimpleDocTemplate(filename, pagesize=A4)  # alternatively use bottomup=False
                    width, height = A4
                    print(f'page width : {width}, height : {height}')

                    styles = getSampleStyleSheet()

                    name = Paragraph('PARA-ELECTRO', style=styles["title"])
                    date = Paragraph(buy_time)
                    title = Paragraph(f'<b>Invoice :</b> #{facture_id}<br/>', style=styles['Italic'])

                    elems = [name, date, title]

                    style = TableStyle([

                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
                        ('GRID', (0, 0), (-1, -2), 1, colors.black),
                        ('GRID', (-2, -1), (-1, -1), 1, colors.black),
                    ])
                    data.append([' ', ' ', ' ', ' ', sum([int(i[-2]) for i in data[1:]] ), f'{sum([int(i[-1].split(" ")[0]) for i in data[1:]] )} DH'])
                    selles_table = Table(data)
                    selles_table.setStyle(style)
                    ttl1 = Paragraph('Sells Operations : ', style=styles["h3"])
                    elems.append(ttl1)
                    elems.append(selles_table)


                    doc.build(elems)
                    notification.notify(title='Invoice Saved Successfully', message=f'The Invoice File (pdf file) saved at {filename}', timeout=5)

                else:
                    notification.notify(title='canceled', message='the facture file has been canceled', timeout=5)

            self.print_facture.setChecked(False)
            # self.operation_type.setChecked(False)
            conn.close()
            [self.to_buy_table.removeRow(0) for _ in range(self.to_buy_table.rowCount())]

            self.product_name.setCurrentIndex(0)
            self.categorie_.setText('')
            self.qt_.setText('')
            self.seller.setText('')
            self.price.setText('')
            self.is_paid.setChecked(True)
            self.f_date.setText(str(strftime('%d-%M-%Y')))
        else:
            notification.notify(title='EROOR', message='Nothing to buy right now, please add some goods .', app_name='PARA-ELECTRO', app_icon=os.path.join(os.path.dirname(__file__), '/src/icons/logo.png'), timeout=8)

    def table_select_event(self):
        # row = [str(i.text()) for i in self.to_buy_table.selectedItems()]
        self.clear.setEnabled(True) if [idx.row() for idx in self.to_buy_table.selectionModel().selectedRows()] else self.clear.setEnabled(False)
        # self.selected_rows = [idx.row() for idx in self.to_buy_table.selectionModel().selectedRows()]
        # if self.selected_rows:
        #     self.clear.setText('Delete')
        #     self.clear.clicked.connect(self.del_row)
        # else:
        #     self.clear.setText('clear')
        #     self.clear.clicked.connect(self.clear_table)

    def del_row(self):
        selected_rows = [idx.row() for idx in self.to_buy_table.selectionModel().selectedRows()]
        # [self.to_buy_table.removeRow(self.to_buy_table.indexAt(int(i)).row()) for i in selected_rows]
        [self.to_buy_table.removeRow(i) for i in selected_rows]



    def add_row(self):

        if self.product_name.currentText() == '' or len(self.product_name.currentText().split( ' - ')) < 2:
            self.err.setText('<font color="red">Invalid Product Name</font>, Hint : code - name...')
        elif self.categorie_.text() == '':
            self.err.setText('<font color="red">Invalid Categorie</font>')
        elif not self.qt_.text() or int(self.qt_.text()) <= 0:
            self.err.setText('<font color="red">Invalid Quantity Value</font>')
        elif len(str(self.price.text()).strip()) == 0 or int(self.price.text()) <= 0 :
            self.err.setText('<font color="red">Invalid Price Value</font>')

        else:
            if self.seller.text() == '':
                self.seller.setText('-')

            if self.product_name.currentText() and self.product_name.currentText().split( ' - ')[0].strip() not in self.all_codes:
                notification.notify(title='error', message='it\'s seems you have no such product with this infos')
                return


            self.err.setText('')
            # self.to_buy_table.addTableRow( [self.product_name.currentText().split(' - ')[1], self.categorie.text(), self.qt.text, self.seller.text(), self.price.text() , 'Yes' if self.paid.isChecked() else 'No'])
            row = self.to_buy_table.rowCount()
            self.to_buy_table.insertRow(0)
            col = 0
            for item in [self.product_name.currentText().split(' - ')[1], self.product_name.currentText().split(' - ')[0], self.categorie_.text(), self.qt_.text(),
                         self.seller.text(), self.price.text(), 'Yes' if self.is_paid.isChecked() else 'No', f'{int(self.qt_.text())*int(self.price.text())} DH']:
                cell = QTableWidgetItem(str(item))
                self.to_buy_table.setItem(0, col, cell)
                col += 1

            self.seller.setText('')
            self.product_name.setCurrentIndex(0)
            self.is_paid.setChecked(True)
            self.qt_.setText('')



    def qt_changing(self):
        if self.qt_.text():
            if str(self.product_name.currentText()).strip() == '':
                self.rest.setText('<font color="red">ERR :</font> Select product first')
                return
            qt = self.qt_value
            code = str(self.product_name.currentText()).split('-')[0].strip()

            for r in range(self.to_buy_table.rowCount()):
                if self.to_buy_table.item(r, 1).text() == code:
                    if str(self.to_buy_table.item(r, 4).text()).lower() == 'buy':
                        qt += int(self.to_buy_table.item(r, 3).text())
                    else:
                        qt -= int(self.to_buy_table.item(r, 3).text())



            if int(qt - int(self.qt_.text()))>= 0:
                self.rest.setText(f'{qt - int(self.qt_.text())}"in the stock')
                self.add.setEnabled(True)
            else:
                self.rest.setText(f'<font color="red">Error</font> ')
                self.add.setEnabled(False)


            # try:
            #     if self.operation_type.isChecked():
            #         self.rest.setText(f'{int(self.qt_.text()) + self.qt_value}"in the stock')
            #     else:
            #         self.rest.setText(f'{int(self.qt_.text()) - self.qt_value}"in the stock')
            # except:
            #     self.rest.setText('<font color="red">ERR :</font> Fill the Quantity')
            # try:
            #     if int(self.qt_.text()) < 1:
            #         self.rest.setText('<font color="red">ERR :</font> Must be > 0')
            #         self.add.setEnabled(False)
            #     else:
            #         self.add.setEnabled(True)
            # except:pass

        else:
            # self.rest.setText('')
            self.add.setEnabled(False)
            self.rest.setText('')

    def product_choosen(self):
        conn = sqlite3.connect(DB)
        curs = conn.cursor()
        code = str(self.product_name.currentText()).split("-")[0].replace(" ", "")
        dt = curs.execute(f'select categorie, qt, sell_price, price from products where code like "{code}"').fetchone()
        if dt:
            self.categorie_.setText(dt[0])
            self.categorie_.setReadOnly(True)
            self.qt_value = int(dt[1])
            self.rest.setText(f'{self.qt_value}"in the stock')

            self.price.setText(str(dt[2]))
            # self.price.setReadOnly(True)
        else:
            self.categorie_.setText('')
            self.categorie_.setReadOnly(False)
            self.rest.setText('0 in the stock')
            self.price.setText('')
            # self.qt_.setText('1')
            self.qt_value = 0
            # self.price.setReadOnly(False)


        conn.close()


    def refresh(self):

        conn = sqlite3.connect(DB)
        self.curs = conn.cursor()
        # ll = list([f' - '.join(i) for i in self.curs.execute(f'select code, name from products where name like "%{self.product_name.text()}%" or code like "%{self.product_name.text()}%"').fetchall()])
        ll = list([f' - '.join(i) for i in self.curs.execute(f'select code, name from products order by name asc').fetchall()])
        ll.insert(0, '')
        # ll = ['yassine', 'baghdadi', 'guercif']
        try:
            self.model.setStringList(list(set([str(i[0]).strip() for i in self.curs.execute('select person from history order by person asc').fetchall()])))
        except Exception as e :
            print(e)
        self.product_name.clear()
        self.product_name.addItems(ll)
        self.all_codes = [i[0] for i in self.curs.execute('select code from products').fetchall()]
        conn.close()
    # def eventFilter(self, source, event) :
    #     # if source is self.product_name or source is self.categorie:
    #     #     if event.type() == QtCore.QEvent.FocusIn:
    #     #         print(f'connecting to {DB}')
    #     #         self.conn = sqlite3.connect(DB)
    #     #         self.curs = self.conn.cursor()
    #     #     elif event.type() == QtCore.QEvent.FocusOut:
    #     #         print(f'deconnecting to {DB}')
    #     #         self.conn.close()
    #
    #     return super(Buy, self).eventFilter(source, event)




#
#
#
# class Buy(QtWidgets.QFrame):
#     def __init__(self):
#         super(Buy, self).__init__()
#         uic.loadUi(os.path.join(os.getcwd(), 'src/ui/buy.ui'), self)
#
#         self.product_name.installEventFilter(self)
#
#         self.completer = QCompleter()
#         self.model = QStringListModel()
#         self.completer.setModel(self.model)
#         # # self.product_name.textChanged.connect(self.product_txt_changing)
#         # self.product_txt_changing()
#         self.seller.installEventFilter(self)
#         self.seller.setCompleter(self.completer)
#         self.product_name.currentTextChanged.connect(self.product_choosen)
#         self.qt_.textChanged.connect(self.qt_changing)
#         self.qt_.installEventFilter(self)
#
#         self.qt_.setValidator(QIntValidator())
#         self.price.setValidator(QIntValidator())
#         self.price.installEventFilter(self)
#         # conn.close()
#         self.qt_value = 0
#         self.add.clicked.connect(self.add_row)
#
#         self.to_buy_table.itemSelectionChanged.connect(self.table_select_event)
#         self.table_header = [ 'Product', 'Code', 'Categorie', 'Quantity', 'Seller', 'Price', 'Paid', 'Total']
#         self.to_buy_table.setColumnCount(len(self.table_header))
#         self.to_buy_table.setHorizontalHeaderLabels(self.table_header)
#         self.to_buy_table.resizeColumnsToContents()
#         # self.to_buy_table.horizontalHeader().setSectionResizeMode(self.table_header.index(self.table_header[-1]), QHeaderView.Stretch)
#         for i in range(len(self.table_header)):
#             self.to_buy_table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)
#
#         self.clear.setEnabled(False)
#         self.clear.clicked.connect(self.del_row)
#         self.buy.clicked.connect(self.save_buy)
#         self.f_date.setText('-'.join(str(datetime.datetime.today().date()).split('-')[::-1]))
#         [self.to_buy_table.removeRow(0) for _ in range(self.to_buy_table.rowCount())]
#
#         self.refresh()
#         # with keyboard.GlobalHotKeys({
#         #     '<ctrl>+<alt>+h': self.save_buy()}) as h:
#         #     h.join()
#
#     def eventFilter(self, obj, event):
#         if event.type() == QtCore.QEvent.KeyPress :
#             if event.key() == QtCore.Qt.Key_Return :
#                 if obj is self.qt_ or obj is self.seller or obj is self.price:
#                     self.add_row()
#
#         return super().eventFilter(obj, event)
#
#     def save_buy(self):
#         if self.to_buy_table.rowCount():
#             # print(self.to_buy_table.rowCount())
#             # print(self.to_buy_table.columnCount())
#             bought = []
#             to_update = []
#             to_add = []
#             conn = sqlite3.connect(DB)
#             curs = conn.cursor()
#             exists_codes = [str(i[0]) for i in curs.execute('select code from products').fetchall()]
#
#
#             exists_invoices_ids = [str(i[0]) for i in curs.execute('SELECT invoice FROM history').fetchall()]
#             pp = 'a b c d e f g h i j k l m n o p q r s t u v w x y z A B C D E F G H J L M N P Q R S T U V W Y X Z 0 1 2 3 4 5 6 7 8 9'
#             def get_facture_id() -> str:
#                 id = ''.join(random.choices(pp.split(), k=10))
#                 if id in exists_invoices_ids:
#                     get_facture_id()
#                 return id
#             facture_id = get_facture_id()
#             for r in range(self.to_buy_table.rowCount()):
#                 if str(self.to_buy_table.item(r, 1).text()) in exists_codes:
#                     tmp = []
#                     for c in range(self.to_buy_table.columnCount()):
#                         tmp.append(str(self.to_buy_table.item(r, c).text()))
#                     to_update.append(tmp)
#                 else:
#                     tmp = []
#                     for c in range(self.to_buy_table.columnCount()):
#                         itm = str(self.to_buy_table.item(r, c).text())
#                         tmp.append(itm)
#                     to_add.append(tmp)
#
#
#
#             buy_time =f'{self.f_date.text()} {str(strftime(f"%H:%M:%S", gmtime()))}' if self.f_date.text() else str(strftime(f"%Y-%m-%d %H:%M:%S", gmtime()))
#             def add_history(i):
#
#                 '''date_time TEXT, product TEXT, qt INTEGER, total_price TEXT, operation TEXT, person TEXT, paid TEXT, invoice TEXT);'''
#
#                 curs.execute(f'''INSERT INTO history (date_time, product, qt, total_price, operation, person, paid, invoice) values(
#                                                         "{buy_time}",
#                                                         "{i[0]}({i[1]})",
#                                                         "{i[3]}",
#                                                         "{int(i[5]) * int(i[3])}",
#                                                         "Buy",
#                                                         "{i[4]}",
#                                                         "{"No" if i[6][0].upper() == "N" else buy_time}",
#                                                         "{facture_id}")''')
#                 conn.commit()
#
#                 # bought.append(curs.execute('select seq from sqlite_sequence where name="history"').fetchone()[0])
#             for i in to_add:
#                 curs.execute(f'''INSERT INTO products (name, code, categorie, qt, price) values("{i[0]}", "{i[1]}", "{i[2]}", {int(i[3])}, "{i[5]}")''')
#                 add_history(i)
#
#             for i in to_update:
#                 curs.execute(f'''
#                      UPDATE products SET qt = qt + {int(i[3])} where code like "{i[1]}"
#                 ''')
#                 add_history(i)
#
#             tt_qt = 0
#             for i in range(self.to_buy_table.rowCount()):
#                 tt_qt += int(self.to_buy_table.item(i, 3).text())
#
#             tt_price = 0
#             for i in range(self.to_buy_table.rowCount()):
#                 tt_price += int(self.to_buy_table.item(i, 5).text())
#
#
#
#             # f_dd = f'{self.f_date.text()} {str(strftime("%H:%M:%S", gmtime()))}' if self.f_date.text() else str(strftime("%Y-%m-%d %H:%M:%S", gmtime()))
#             # curs.execute(f"""insert into factures (date_time, persons, products, qt, total_price, operation) values(
#             #         "{f_dd}",
#             #         "{'-'.join(list(str(self.to_buy_table.item(i, 4).text()) for i in range(self.to_buy_table.rowCount())))}",
#             #         "{'-'.join([str(i) for i in bought])}",
#             #         {tt_qt},
#             #         "{str(tt_price)}",
#             #         "buy"
#             # )""")
#             debting = []
#
#             # for r in range(self.to_buy_table.rowCount()):
#             #     tt = []
#             #     if self.to_buy_table.item(r, 6).text()[0].upper() == 'N':
#             #         for c in range(self.to_buy_table.columnCount()):
#             #             tt.append(self.to_buy_table.item(r, c).text())
#             #     debting.append(tt)
#             #
#             # print(enumerate(debting))
#             # for i , v in enumerate(debting):
#
#             # curs.execute(f'''
#             #     insert into debt (date_time, person, )
#             # ''')
#             # print(debting)
#             conn.commit()
#
#             # print(f'rows : {self.to_buy_table.rowCount()} , Clolumns : {self.to_buy_table.columnCount()}')
#             # self.to_buy_table.clear()
#             # self.to_buy_table.setColumnCount(len(self.table_header))
#             # self.to_buy_table.setHorizontalHeaderLabels(self.table_header)
#
#             if self.print_facture.isChecked():
#
#                 # doc = Document()
#                 # doc.add_heading(STORE_NAME, 0)
#                 # # f_name = f'''Facture N°: {curs.execute("select seq from sqlite_sequence where name='factures'").fetchone()[0]}, for : {f_dd.replace(' ', '_') if platform.system() == 'Linux' else f_dd.replace(':', '-').replace(' ', '_')}'''
#                 # f_name = f'''Facture Id: {facture_id}, for : {buy_time}'''
#                 #
#                 # doc.add_heading(f_name, 1)
#                 #
#                 # table = doc.add_table(rows=self.to_buy_table.rowCount(), cols=self.to_buy_table.columnCount())
#                 # table.allow_autofit = True
#                 # table.alignment = WD_TABLE_ALIGNMENT.LEFT
#                 # table.style = 'Table Grid'
#                 # sections = doc.sections
#                 # for section in sections:
#                 #     section.top_margin = Cm(1.5)
#                 #     section.bottom_margin = Cm(1.5)
#                 #     section.left_margin = Cm(1.5)
#                 #     section.right_margin = Cm(1.5)
#                 #
#                 #
#                 # hdr_cells = table.rows[0].cells
#                 # hdr_labels = [i for i in 'Product Categorie Quantity Seller is_Paid Total'.split(' ')]
#                 #
#                 # for c in range(len(hdr_labels)):
#                 #     # cc = self.to_buy_table.horizontalHeaderItem(c).text()
#                 #     hdr_cells[c].text = str(hdr_labels[c])
#                 #
#                 # rows = []
#                 # table.left_margin = Cm(1.5)
#                 # table.right_margin = Cm(1.5)
#                 # for row in range(self.to_buy_table.rowCount()):
#                 #     rr = []
#                 #     for col in [0, 1, 2, 3, 4, 6, 7]:
#                 #         rr.append(str(self.to_buy_table.item(row, col).text()))
#                 #     rows.append(rr)
#                 # # print(rows)
#                 # for row in rows:
#                 #     tr = table.add_row()
#                 #     for idx, col in enumerate(row):
#                 #         tr.cells[idx].text = col if col else '---'
#
#                 qts = []
#                 prices = []
#                 rows = [[i.replace('_', ' ') for i in 'Product Categorie Seller is_Paid Quantity Total'.split(' ')]]
#                 for row in range(self.to_buy_table.rowCount()):
#                     rr = [f'{str(self.to_buy_table.item(row, 0).text())}({str(self.to_buy_table.item(row, 1).text())})']
#                     for col in [2, 4, 6, 3, 7]:
#                         rr.append(str(self.to_buy_table.item(row, col).text()))
#                     qts.append(int(str(self.to_buy_table.item(row, 3).text())))
#                     prices.append(int(str(self.to_buy_table.item(row, 7).text()).split(' ')[0]))
#                     rows.append(rr)
#
#                 rows.append([' ', ' ', ' ', ' ', sum(qts), f'{sum(prices)} DH'])
#
#                 if filename := QtWidgets.QFileDialog.getSaveFileName(caption='Print Facture', filter="PDF (*.pdf )", directory=os.path.join(DESKTOP, f'invoice_{facture_id}'))[0]:
#                     if not QFileInfo(filename).suffix():
#                         filename += '.pdf'
#
#                     # doc = SimpleDocTemplate(filename=filename, pagesize=letter, bottomMargin=18)
#                     # table = Table(data=rows)
#                     doc = SimpleDocTemplate(filename, pagesize=A4)  # alternatively use bottomup=False
#                     width, height = A4
#                     print(f'page width : {width}, height : {height}')
#
#
#
#
#                     table = Table(rows)
#                     style = TableStyle([
#
#                         ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
#                         ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
#                         ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
#                         ('BACKGROUND', (0, 1), (-1, -2), colors.beige),
#                         ('GRID', (0, 0), (-1, -2), 1, colors.black),
#                         ('GRID', (-2, -1), (-1, -1), 1, colors.black),
#                     ])
#
#                     table.setStyle(style)
#
#
#                     styles = getSampleStyleSheet()
#                     name = Paragraph('PARA-ELECTRO', style=styles["title"])
#                     date = Paragraph(buy_time)
#                     title = Paragraph(f'<b>Invoice :</b> #{facture_id}<br/>', style=styles['Italic'])
#
#                     doc.build([name, date, title, table])
#
#                 else:
#                     notification.notify(title='canceled', message='the facture file has been canceled', timeout=5)
#
#             self.print_facture.setChecked(False)
#             conn.close()
#             [self.to_buy_table.removeRow(0) for _ in range(self.to_buy_table.rowCount())]
#
#             self.product_name.setCurrentIndex(0)
#             self.categorie_.setText('')
#             self.qt_.setText('')
#             self.seller.setText('')
#             self.price.setText('')
#             self.is_paid.setChecked(True)
#             self.f_date.setText(str(strftime('%d-%M-%Y')))
#         else:
#             notification.notify(title='EROOR', message='Nothing to buy right now, please add some goods .', app_name='PARA-ELECTRO', app_icon=os.path.join(os.path.dirname(__file__), '/src/icons/logo.png'), timeout=8)
#
#     def table_select_event(self):
#         # row = [str(i.text()) for i in self.to_buy_table.selectedItems()]
#         self.clear.setEnabled(True) if [idx.row() for idx in self.to_buy_table.selectionModel().selectedRows()] else self.clear.setEnabled(False)
#         # self.selected_rows = [idx.row() for idx in self.to_buy_table.selectionModel().selectedRows()]
#         # if self.selected_rows:
#         #     self.clear.setText('Delete')
#         #     self.clear.clicked.connect(self.del_row)
#         # else:
#         #     self.clear.setText('clear')
#         #     self.clear.clicked.connect(self.clear_table)
#
#     def del_row(self):
#         selected_rows = [idx.row() for idx in self.to_buy_table.selectionModel().selectedRows()]
#         # [self.to_buy_table.removeRow(self.to_buy_table.indexAt(int(i)).row()) for i in selected_rows]
#         [self.to_buy_table.removeRow(i) for i in selected_rows]
#
#
#
#     def add_row(self):
#         if self.product_name.currentText() == '' or len(self.product_name.currentText().split( ' - ')) < 2:
#             self.err.setText('<font color="red">Invalid Product Name</font>, Hint : code - name...')
#         elif self.categorie_.text() == '':
#             self.err.setText('<font color="red">Invalid Categorie</font>')
#         elif not self.qt_.text() or int(self.qt_.text()) <= 0:
#             self.err.setText('<font color="red">Invalid Quantity Value</font>')
#         elif self.seller.text() == '':
#             self.err.setText('<font color="red">Invalid Seller Name</font>, Hint : write - if you don\'t have one')
#         elif len(str(self.price.text()).strip()) == 0 or int(self.price.text()) <= 0 :
#             self.err.setText('<font color="red">Invalid Price Value</font>')
#         else:
#             self.err.setText('')
#             # self.to_buy_table.addTableRow( [self.product_name.currentText().split(' - ')[1], self.categorie.text(), self.qt.text, self.seller.text(), self.price.text() , 'Yes' if self.paid.isChecked() else 'No'])
#             row = self.to_buy_table.rowCount()
#             self.to_buy_table.insertRow(0)
#             col = 0
#             for item in [self.product_name.currentText().split(' - ')[1], self.product_name.currentText().split(' - ')[0], self.categorie_.text(), self.qt_.text(),
#                          self.seller.text(), self.price.text(), 'Yes' if self.is_paid.isChecked() else 'No', f'{int(self.qt_.text())*int(self.price.text())} DH']:
#                 cell = QTableWidgetItem(str(item))
#                 self.to_buy_table.setItem(0, col, cell)
#                 col += 1
#
#             self.seller.setText('')
#             self.product_name.setCurrentIndex(0)
#             self.is_paid.setChecked(True)
#             self.qt_.setText('')
#
#     def qt_changing(self):
#         if self.qt_.text():
#             if self.product_name.currentText() == '':
#                 self.rest.setText('<font color="red">ERR :</font> Select product first')
#                 return
#
#             try:
#                 # old = int(self.rest.text().split('*')[0])
#
#                 self.rest.setText(f'{int(self.qt_.text())+ self.qt_value}"in the stock')
#             except:
#                 self.rest.setText('<font color="red">ERR :</font> Fill the Quantity')
#             try:
#                 if int(self.qt_.text()) < 1:
#                     self.rest.setText('<font color="red">ERR :</font> Must be > 0')
#                     self.add.setEnabled(False)
#                 else:self.add.setEnabled(True)
#             except:pass
#         else:
#             # self.rest.setText('')
#             self.add.setEnabled(False)
#
#     def product_choosen(self):
#         conn = sqlite3.connect(DB)
#         curs = conn.cursor()
#         code = str(self.product_name.currentText()).split("-")[0].replace(" ", "")
#         # print(code)
#         try:
#             dt = curs.execute(f'select categorie, qt, price from products where code like "{code}"').fetchone()
#             self.categorie_.setText(dt[0])
#             self.categorie_.setReadOnly(True)
#             self.qt_value = int(dt[1])
#             self.rest.setText(f'{self.qt_value}"in the stock')
#
#             self.price.setText(str(dt[2]))
#
#         except:
#             self.categorie_.setText('')
#             self.categorie_.setReadOnly(False)
#             self.rest.setText('0 in the stock')
#             self.price.setText('')
#             # self.qt_.setText('1')
#             self.qt_value = 0
#
#
#     def refresh(self):
#
#         conn = sqlite3.connect(DB)
#         self.curs = conn.cursor()
#         # ll = list([f' - '.join(i) for i in self.curs.execute(f'select code, name from products where name like "%{self.product_name.text()}%" or code like "%{self.product_name.text()}%"').fetchall()])
#         ll = list([f' - '.join(i) for i in self.curs.execute(f'select code, name from products order by name asc').fetchall()])
#         ll.insert(0, '')
#         # ll = ['yassine', 'baghdadi', 'guercif']
#         try:
#             self.model.setStringList(list(set([str(i[0]).strip() for i in self.curs.execute('select person from history order by person asc').fetchall()])))
#         except Exception as e :
#             print(e)
#         self.product_name.clear()
#         self.product_name.addItems(ll)
#         conn.close()
#     # def eventFilter(self, source, event) :
#     #     # if source is self.product_name or source is self.categorie:
#     #     #     if event.type() == QtCore.QEvent.FocusIn:
#     #     #         print(f'connecting to {DB}')
#     #     #         self.conn = sqlite3.connect(DB)
#     #     #         self.curs = self.conn.cursor()
#     #     #     elif event.type() == QtCore.QEvent.FocusOut:
#     #     #         print(f'deconnecting to {DB}')
#     #     #         self.conn.close()
#     #
#     #     return super(Buy, self).eventFilter(source, event)
#





class Debt(QtWidgets.QDialog):
    def __init__(self):
        super(Debt, self).__init__()
        uic.loadUi(os.path.join(os.getcwd(), 'src/ui/debt.ui'), self)




if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setWindowIcon(QtGui.QIcon('src/icons/logo.png'))
    main = Main()
    main.show()
    sys.exit(app.exec_())